from pathlib import Path
from typing import Dict, List, Optional, Union
import logging
from datetime import datetime
import os
import json
from docx import Document
import subprocess

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Class for generating quote documents."""
    
    def __init__(self, template_dir: str = "templates", output_dir: str = "output"):
        """Initialize document generator."""
        self.template_dir = Path(template_dir)
        self.output_dir = Path(output_dir)
        self._ensure_directories_exist()
    
    def _ensure_directories_exist(self):
        """Ensure template and output directories exist."""
        self.template_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        for quote_type in ["auto", "home", "specialty"]:
            (self.output_dir / quote_type).mkdir(parents=True, exist_ok=True)
    
    def generate_quote_documents(self, quote_data: dict) -> Dict[str, str]:
        """Generate documents for all quote types in quote_data."""
        logger.info(f"Starting document generation for quote {quote_data.get('id')}")
        documents = {}
        
        try:
            # Generate documents based on quote types
            quote_types = quote_data.get("quote_types", [])
            client_name = quote_data.get("client_name", "Unknown")
            
            if "AUTO" in quote_types:
                logger.info("Generating auto quote document")
                docx_path, pdf_path = self._generate_auto_quote(quote_data, client_name)
                documents["auto"] = {
                    "docx_path": docx_path,
                    "pdf_path": pdf_path
                }
            
            if "HOME" in quote_types:
                logger.info("Generating home quote document")
                docx_path, pdf_path = self._generate_home_quote(quote_data, client_name)
                documents["home"] = {
                    "docx_path": docx_path,
                    "pdf_path": pdf_path
                }
            
            if "SPECIALTY" in quote_types:
                logger.info("Generating specialty quote document")
                docx_path, pdf_path = self._generate_specialty_quote(quote_data, client_name)
                documents["specialty"] = {
                    "docx_path": docx_path,
                    "pdf_path": pdf_path
                }
            
            logger.info(f"Document generation completed for quote {quote_data.get('id')}")
            return documents
            
        except Exception as e:
            logger.error(f"Error in document generation: {str(e)}")
            raise
    
    def _generate_auto_quote(self, quote_data: dict, client_name: str) -> tuple:
        """Generate auto quote document."""
        try:
            template_path = self.template_dir / "auto_quote_template.docx"
            if not template_path.exists():
                logger.error(f"Auto quote template not found: {template_path}")
                return None, None
                
            doc = Document(template_path)
            
            # Prepare personal info data
            personal_info = quote_data.get("personal_info", {})
            
            # Prepare replacement dictionary
            replacements = {
                "client_name": client_name,
                "client_email": quote_data.get("client_email", "N/A"),
                "client_phone": quote_data.get("client_phone", "N/A"),
                "address": quote_data.get("address", "N/A"),
                "created_at": quote_data.get("created_at", "N/A"),
                "effective_date": quote_data.get("effective_date", "N/A")
            }
            
            # Add auto-specific data if available
            auto_data = quote_data.get("auto_data", {})
            if auto_data:
                replacements.update({
                    "current_carrier": auto_data.get("current_carrier", "N/A"),
                    "years_with_carrier": str(auto_data.get("years_with_carrier", "N/A")),
                    "expiration_date": auto_data.get("expiration_date", "N/A"),
                    "current_limits": auto_data.get("current_limits", "N/A"),
                    "quoting_limits": auto_data.get("quoting_limits", "N/A")
                })
            
            # Replace placeholders in document
            self._replace_placeholders(doc, replacements)
            
            # Save document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            sanitized_name = "".join(x for x in client_name if x.isalnum() or x.isspace()).replace(" ", "_")
            output_filename = f"auto_{sanitized_name}_{timestamp}.docx"
            output_path = self.output_dir / "auto" / output_filename
            logger.info(f"Saving document to: {output_path}")
            doc.save(str(output_path))
            
            # Convert to PDF
            pdf_path = output_path.with_suffix(".pdf")
            try:
                self._convert_to_pdf(str(output_path), str(pdf_path))
                logger.info(f"PDF saved to: {pdf_path}")
            except Exception as e:
                logger.warning(f"PDF conversion failed: {str(e)}")
                pdf_path = None
            
            return str(output_path), str(pdf_path) if pdf_path else None
            
        except Exception as e:
            logger.error(f"Error generating auto quote document: {str(e)}")
            raise
    
    def _generate_home_quote(self, quote_data: dict, client_name: str) -> tuple:
        """Generate home quote document."""
        try:
            template_path = self.template_dir / "home_quote_template.docx"
            if not template_path.exists():
                logger.error(f"Home quote template not found: {template_path}")
                return None, None
                
            doc = Document(template_path)
            
            # Prepare replacement dictionary
            replacements = {
                "client_name": client_name,
                "client_email": quote_data.get("client_email", "N/A"),
                "client_phone": quote_data.get("client_phone", "N/A"),
                "address": quote_data.get("address", "N/A"),
                "created_at": quote_data.get("created_at", "N/A"),
                "effective_date": quote_data.get("effective_date", "N/A")
            }
            
            # Add home details if available
            home_details = quote_data.get("home_details", {})
            if home_details:
                replacements.update({
                    "year_built": home_details.get("year_built", "N/A"),
                    "square_footage": home_details.get("square_footage", "N/A"),
                    "construction_type": home_details.get("construction_type", "N/A"),
                    "roof_type": home_details.get("roof_type", "N/A"),
                    "number_of_stories": home_details.get("number_of_stories", "N/A"),
                    "garage_type": home_details.get("garage_type", "N/A"),
                    "basement_type": home_details.get("basement_type", "N/A"),
                    "security_system": "Yes" if home_details.get("security_system") else "No",
                    "swimming_pool": "Yes" if home_details.get("swimming_pool") else "No"
                })
            
            # Replace placeholders in document
            self._replace_placeholders(doc, replacements)
            
            # Save document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            sanitized_name = "".join(x for x in client_name if x.isalnum() or x.isspace()).replace(" ", "_")
            output_filename = f"home_{sanitized_name}_{timestamp}.docx"
            output_path = self.output_dir / "home" / output_filename
            logger.info(f"Saving document to: {output_path}")
            doc.save(str(output_path))
            
            # Convert to PDF
            pdf_path = output_path.with_suffix(".pdf")
            try:
                self._convert_to_pdf(str(output_path), str(pdf_path))
                logger.info(f"PDF saved to: {pdf_path}")
            except Exception as e:
                logger.warning(f"PDF conversion failed: {str(e)}")
                pdf_path = None
            
            return str(output_path), str(pdf_path) if pdf_path else None
            
        except Exception as e:
            logger.error(f"Error generating home quote document: {str(e)}")
            raise
    
    def _generate_specialty_quote(self, quote_data: dict, client_name: str) -> tuple:
        """Generate specialty quote document."""
        try:
            template_path = self.template_dir / "specialty_quote_template.docx"
            if not template_path.exists():
                logger.error(f"Specialty quote template not found: {template_path}")
                return None, None
                
            doc = Document(template_path)
            
            # Prepare replacement dictionary
            replacements = {
                "client_name": client_name,
                "client_email": quote_data.get("client_email", "N/A"),
                "client_phone": quote_data.get("client_phone", "N/A"),
                "address": quote_data.get("address", "N/A"),
                "created_at": quote_data.get("created_at", "N/A"),
                "effective_date": quote_data.get("effective_date", "N/A")
            }
            
            # Add specialty items if available
            specialty_items = quote_data.get("specialty_items", [])
            for i, item in enumerate(specialty_items[:8], 1):  # Limit to 8 items
                item_prefix = f"item{i}_"
                replacements.update({
                    f"{item_prefix}type": item.get("type", "N/A"),
                    f"{item_prefix}year": item.get("year", "N/A"),
                    f"{item_prefix}make": item.get("make", "N/A"),
                    f"{item_prefix}model": item.get("model", "N/A"),
                    f"{item_prefix}vin": item.get("vin", "N/A"),
                    f"{item_prefix}horsepower": item.get("horsepower", "N/A"),
                    f"{item_prefix}top_speed": item.get("top_speed", "N/A"),
                    f"{item_prefix}market_value": item.get("market_value", "N/A"),
                    f"{item_prefix}storage_location": item.get("storage_location", "N/A"),
                    f"{item_prefix}comp_deductible": item.get("comp_deductible", "N/A"),
                    f"{item_prefix}coll_deductible": item.get("coll_deductible", "N/A")
                })
            
            # Replace placeholders in document
            self._replace_placeholders(doc, replacements)
            
            # Save document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            sanitized_name = "".join(x for x in client_name if x.isalnum() or x.isspace()).replace(" ", "_")
            output_filename = f"specialty_{sanitized_name}_{timestamp}.docx"
            output_path = self.output_dir / "specialty" / output_filename
            logger.info(f"Saving document to: {output_path}")
            doc.save(str(output_path))
            
            # Convert to PDF
            pdf_path = output_path.with_suffix(".pdf")
            try:
                self._convert_to_pdf(str(output_path), str(pdf_path))
                logger.info(f"PDF saved to: {pdf_path}")
            except Exception as e:
                logger.warning(f"PDF conversion failed: {str(e)}")
                pdf_path = None
            
            return str(output_path), str(pdf_path) if pdf_path else None
            
        except Exception as e:
            logger.error(f"Error generating specialty quote document: {str(e)}")
            raise
    
    def _replace_placeholders(self, doc: Document, replacements: Dict[str, str]):
        """Replace placeholders in document with actual values."""
        logger.info("Replacing placeholders in document")
        logger.debug(f"Data keys available: {list(replacements.keys())}")
        
        for paragraph in doc.paragraphs:
            placeholder_keys = []
            for key in replacements.keys():
                if f"{{{{{key}}}}}" in paragraph.text:
                    placeholder_keys.append(key)
            
            if placeholder_keys:
                logger.debug(f"Found placeholders in paragraph: {placeholder_keys}")
                for key in placeholder_keys:
                    value = replacements[key]
                    placeholder = f"{{{{{key}}}}}"
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
                    logger.debug(f"Replaced '{key}' with '{value}'")
                logger.debug(f"Paragraph changed from '{paragraph.text}' to '{paragraph.text}'")
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        placeholder_keys = []
                        for key in replacements.keys():
                            if f"{{{{{key}}}}}" in paragraph.text:
                                placeholder_keys.append(key)
                        
                        if placeholder_keys:
                            for key in placeholder_keys:
                                value = replacements[key]
                                placeholder = f"{{{{{key}}}}}"
                                paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        logger.info("Placeholder replacement completed successfully")

    def _convert_to_pdf(self, docx_path: str, pdf_path: str):
        """Convert DOCX to PDF using docx2pdf or LibreOffice as fallback."""
        try:
            # First try with docx2pdf (requires MS Word)
            from docx2pdf import convert
            convert(docx_path, pdf_path)
        except Exception as e:
            logger.warning(f"docx2pdf conversion failed: docx2pdf is not implemented for linux as it requires Microsoft Word to be installed, trying LibreOffice")
            
            # Fallback to LibreOffice
            cmd = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                   os.path.dirname(pdf_path), docx_path]
            
            try:
                subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                logger.info(f"LibreOffice PDF conversion successful: {pdf_path}")
            except subprocess.CalledProcessError as e:
                logger.error(f"LibreOffice conversion failed: {e}")
                raise
            except FileNotFoundError:
                logger.error("LibreOffice not found, cannot convert to PDF")
                raise

def generate_quote_documents(quote_data: Dict, template_dir: Optional[Path] = None, output_dir: Optional[Path] = None) -> Dict[str, str]:
    """Generate documents for a quote request."""
    template_dir = template_dir or os.getenv("TEMPLATE_DIR", "templates")
    output_dir = output_dir or os.getenv("OUTPUT_DIR", "output")
    
    generator = DocumentGenerator(template_dir, output_dir)
    return generator.generate_quote_documents(quote_data) 