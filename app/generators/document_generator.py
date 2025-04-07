from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import os
from datetime import datetime
from pathlib import Path

class DocumentGenerator:
    def __init__(self, output_dir="output"):
        self.output_dir = Path(output_dir)
        self.templates_dir = Path("templates")
        self.output_dir.mkdir(exist_ok=True)
        self.templates_dir.mkdir(exist_ok=True)

    def _create_output_filename(self, quote_type, client_name):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = "".join(c if c.isalnum() else "_" for c in client_name)
        return f"{quote_type}_{safe_name}_{timestamp}"

    def _replace_placeholders(self, doc, data):
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if f"{{{{{key}}}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value or "N/A"))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in data.items():
                            if f"{{{{{key}}}}}" in paragraph.text:
                                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value or "N/A"))

    def generate_documents(self, quote_type, data):
        # Create filenames
        base_filename = self._create_output_filename(quote_type, data.get("primary_insured_name", "quote"))
        docx_path = self.output_dir / f"{base_filename}.docx"
        pdf_path = self.output_dir / f"{base_filename}.pdf"

        # Load template
        template_path = self.templates_dir / f"{quote_type}_template.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        doc = Document(template_path)
        
        # Replace placeholders
        self._replace_placeholders(doc, data)

        # Save DOCX
        doc.save(docx_path)

        # Convert to PDF
        convert(docx_path, pdf_path)

        return {
            "docx": str(docx_path),
            "pdf": str(pdf_path)
        }

    def generate_multiple_quotes(self, quote_types, data):
        results = {}
        for quote_type in quote_types:
            results[quote_type] = self.generate_documents(quote_type, data)
        return results 