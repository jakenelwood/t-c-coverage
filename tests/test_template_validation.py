import unittest
from pathlib import Path
import docx
import os

class TestTemplateValidation(unittest.TestCase):
    def setUp(self):
        self.templates_dir = Path("app/templates")
        self.required_placeholders = {
            "auto": [
                "{{name}}", "{{email}}", "{{phone}}", "{{address}}",
                "{{dob}}", "{{ssn}}", "{{marital_status}}", "{{occupation}}",
                "{{vehicle_year}}", "{{vehicle_make}}", "{{vehicle_model}}",
                "{{vehicle_vin}}", "{{vehicle_usage}}", "{{vehicle_miles}}",
                "{{vehicle_driver}}", "{{comp_deductible}}", "{{coll_deductible}}",
                "{{glass_coverage}}", "{{towing}}", "{{rental}}", "{{financed}}",
                "{{gap}}"
            ],
            "home": [
                "{{name}}", "{{email}}", "{{phone}}", "{{address}}",
                "{{dob}}", "{{ssn}}", "{{marital_status}}", "{{occupation}}",
                "{{year_built}}", "{{square_footage}}", "{{construction_type}}",
                "{{roof_type}}", "{{stories}}", "{{garage_type}}",
                "{{basement_type}}", "{{dwelling_coverage}}", "{{personal_property}}",
                "{{liability}}", "{{medical_payments}}", "{{deductible}}"
            ],
            "specialty": [
                "{{name}}", "{{email}}", "{{phone}}", "{{address}}",
                "{{dob}}", "{{ssn}}", "{{marital_status}}", "{{occupation}}",
                "{{item_type}}", "{{item_year}}", "{{item_make}}",
                "{{item_model}}", "{{item_vin}}", "{{item_hp}}",
                "{{item_speed}}", "{{item_value}}", "{{storage_location}}",
                "{{comp_deductible}}", "{{coll_deductible}}"
            ]
        }

    def test_template_files_exist(self):
        """Test that all required template files exist."""
        self.assertTrue(self.templates_dir.exists())
        self.assertTrue((self.templates_dir / "auto-request-frm.docx").exists())
        self.assertTrue((self.templates_dir / "home-quote-request-form.docx").exists())
        self.assertTrue((self.templates_dir / "specialty-quote-request-form.docx").exists())

    def test_auto_template_placeholders(self):
        """Test that auto template contains all required placeholders."""
        doc = docx.Document(self.templates_dir / "auto-request-frm.docx")
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        for placeholder in self.required_placeholders["auto"]:
            self.assertIn(placeholder, text, f"Missing placeholder: {placeholder}")

    def test_home_template_placeholders(self):
        """Test that home template contains all required placeholders."""
        doc = docx.Document(self.templates_dir / "home-quote-request-form.docx")
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        for placeholder in self.required_placeholders["home"]:
            self.assertIn(placeholder, text, f"Missing placeholder: {placeholder}")

    def test_specialty_template_placeholders(self):
        """Test that specialty template contains all required placeholders."""
        doc = docx.Document(self.templates_dir / "specialty-quote-request-form.docx")
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        for placeholder in self.required_placeholders["specialty"]:
            self.assertIn(placeholder, text, f"Missing placeholder: {placeholder}")

    def test_template_formatting(self):
        """Test that templates maintain proper formatting."""
        for template_file in self.templates_dir.glob("*.docx"):
            doc = docx.Document(template_file)
            
            # Check font size consistency
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    self.assertEqual(run.font.size.pt, 11, 
                                   f"Font size should be 11pt in {template_file.name}")

            # Check margins
            sections = doc.sections
            for section in sections:
                self.assertEqual(section.left_margin.inches, 1.0)
                self.assertEqual(section.right_margin.inches, 1.0)
                self.assertEqual(section.top_margin.inches, 1.0)
                self.assertEqual(section.bottom_margin.inches, 1.0)

    def test_template_structure(self):
        """Test that templates maintain proper structure."""
        for template_file in self.templates_dir.glob("*.docx"):
            doc = docx.Document(template_file)
            
            # Check for required sections
            has_header = False
            has_footer = False
            has_content = False
            
            for paragraph in doc.paragraphs:
                if "Twin Cities Coverage" in paragraph.text:
                    has_header = True
                if "Generated on" in paragraph.text:
                    has_footer = True
                if any(placeholder in paragraph.text for placeholder in 
                      self.required_placeholders[template_file.stem.split("-")[0]]):
                    has_content = True
            
            self.assertTrue(has_header, f"Missing header in {template_file.name}")
            self.assertTrue(has_footer, f"Missing footer in {template_file.name}")
            self.assertTrue(has_content, f"Missing content in {template_file.name}")

    def test_template_compatibility(self):
        """Test that templates are compatible with docx2pdf conversion."""
        for template_file in self.templates_dir.glob("*.docx"):
            # Check for unsupported elements
            doc = docx.Document(template_file)
            
            # No tables with merged cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.assertFalse(cell._tc.get_or_add_tcPr().xpath('.//w:vMerge'),
                                       f"Merged cells found in {template_file.name}")
            
            # No complex formatting
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    self.assertFalse(run.font.highlight_color,
                                   f"Highlighted text found in {template_file.name}")
                    self.assertFalse(run.font.strike,
                                   f"Strikethrough text found in {template_file.name}")

    def test_template_metadata(self):
        """Test that templates have correct metadata."""
        for template_file in self.templates_dir.glob("*.docx"):
            doc = docx.Document(template_file)
            
            # Check core properties
            core_properties = doc.core_properties
            self.assertEqual(core_properties.author, "Twin Cities Coverage")
            self.assertEqual(core_properties.company, "Twin Cities Coverage")
            self.assertIn("Quote Request Form", core_properties.title)

    def test_template_version_control(self):
        """Test that templates maintain version information."""
        for template_file in self.templates_dir.glob("*.docx"):
            doc = docx.Document(template_file)
            
            # Check for version information in footer
            has_version = False
            for paragraph in doc.paragraphs:
                if "Version" in paragraph.text:
                    has_version = True
                    version_text = paragraph.text
                    self.assertRegex(version_text, r"Version \d+\.\d+\.\d+")
            
            self.assertTrue(has_version, f"Missing version information in {template_file.name}")

if __name__ == '__main__':
    unittest.main() 